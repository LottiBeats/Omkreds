"""
word_builder.py — converts v2 document blocks → Word (.docx) bytes

Block types handled
───────────────────
Document-level (project["documents"][doc_id]["blocks"]):
  heading      → Word Heading 1 / 2 / 3
  text         → Normal paragraph
  image        → embedded picture (decoded from base64 data URL)
  python_calc  → Heading 2 + output text
  beam_fem     → calc block + matplotlib figure
  Any calc type → flatten _result list and render inner calc_core blocks

Calc-core blocks (inside _result):
  module_header → Heading 1, teal/brown/etc. left-bar via run colour
  section       → Heading 2
  h1            → Heading 1
  text          → Normal paragraph
  note          → Normal, italic, amber colour
  calc_row      → batched into a borderless 3-column table
  handcalc      → LaTeX stripped to plain text, shown as calc_row table
  check         → green ✓ or red ✗ coloured line
  table         → Word table with header row
  figure        → embedded picture from file path

Usage
─────
from word_builder import build_word
docx_bytes = build_word(project, blocks, doc_id="A2")
"""

import base64
import io
import os
import re
import tempfile
from typing import List

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Colour palette ─────────────────────────────────────────────────────────────

_GREEN  = RGBColor(0x16, 0xa3, 0x4a)   # check pass
_RED    = RGBColor(0xdc, 0x26, 0x26)   # check fail
_AMBER  = RGBColor(0xb4, 0x53, 0x09)   # note
_BRAND  = RGBColor(0x1e, 0x3a, 0x5f)   # dark-blue heading

_MATERIAL_COLOURS = {
    "steel":    RGBColor(0x12, 0x78, 0x8e),   # teal
    "timber":   RGBColor(0xae, 0x34, 0x19),   # dark orange-brown
    "concrete": RGBColor(0x59, 0x5f, 0x61),   # mid grey
    "masonry":  RGBColor(0x73, 0x1f, 0x0d),   # deep brown
    "general":  RGBColor(0xe7, 0x48, 0x25),   # OMKREDS orange
}

_HEADER_SHADE = "F0F4F8"   # very light blue-grey for calc_row header


# ── LaTeX → plain text (minimal, handles the common cases) ────────────────────

def _latex_to_lines(latex: str) -> List[str]:
    """Strip LaTeX markup and return a list of readable lines."""
    # Remove environments
    body = re.sub(r'\\begin\{[^}]+\}|\\end\{[^}]+\}', '', latex)
    body = re.sub(r'(?<!\\)\\\[|(?<!\\)\\\]', '', body)

    # Split on \\ (line break) ignoring optional spacing
    raw = [p.strip() for p in re.split(r'\\\\\s*(?:\[[\w.]+\])?', body) if p.strip()]

    def clean(s):
        s = re.sub(r'\\text\w*\{([^}]+)\}', '', s)
        s = re.sub(r'\\mathrm\{([^}]+)\}', r'\1', s)
        s = re.sub(r'\\[hv]space\*?\{[^}]+\}', '', s)
        s = re.sub(r'\\dfrac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', s)
        s = re.sub(r'\\frac\{([^}]+)\}\{([^}]+)\}', r'(\1)/(\2)', s)
        s = re.sub(r'\\sqrt\{([^}]+)\}', r'√(\1)', s)
        s = re.sub(r'_\{([^}]+)\}', r'_\1', s)
        s = re.sub(r'\^\{([^}]+)\}', r'^\1', s)
        s = re.sub(r'\{([^}]+)\}', r'\1', s)
        s = re.sub(r'[{}]', '', s)
        s = re.sub(r'\\[,;:!]', '', s)
        s = re.sub(r'\\(?:quad|qquad)\b', '  ', s)
        s = re.sub(r'\\cdot\b', '·', s)
        s = re.sub(r'\\times\b', '×', s)
        s = re.sub(r'\\leq\b', '≤', s)
        s = re.sub(r'\\geq\b', '≥', s)
        s = re.sub(r'\\neq\b', '≠', s)
        s = re.sub(r'\\approx\b', '≈', s)
        s = re.sub(r'\\pm\b', '±', s)
        s = re.sub(r'\\sigma\b', 'σ', s)
        s = re.sub(r'\\tau\b', 'τ', s)
        s = re.sub(r'\\lambda\b', 'λ', s)
        s = re.sub(r'\\omega\b', 'ω', s)
        s = re.sub(r'\\alpha\b', 'α', s)
        s = re.sub(r'\\beta\b', 'β', s)
        s = re.sub(r'\\gamma\b', 'γ', s)
        s = re.sub(r'\\delta\b', 'δ', s)
        s = re.sub(r'\\eta\b', 'η', s)
        s = re.sub(r'\\phi\b', 'φ', s)
        s = re.sub(r'\\chi\b', 'χ', s)
        s = re.sub(r'\\rho\b', 'ρ', s)
        s = re.sub(r'\\pi\b', 'π', s)
        s = re.sub(r'\\mu\b', 'μ', s)
        s = re.sub(r'\\nu\b', 'ν', s)
        s = re.sub(r'\\epsilon\b', 'ε', s)
        s = re.sub(r'\\varepsilon\b', 'ε', s)
        s = re.sub(r'\\varphi\b', 'φ', s)
        s = re.sub(r'\\Sigma\b', 'Σ', s)
        s = re.sub(r'\\Delta\b', 'Δ', s)
        s = re.sub(r'\\Phi\b', 'Φ', s)
        s = re.sub(r'\\infty\b', '∞', s)
        s = re.sub(r'\\rightarrow\b', '→', s)
        s = re.sub(r'\\leftarrow\b', '←', s)
        s = re.sub(r'\\left|\\right', '', s)
        s = re.sub(r'\\[a-zA-Z]+', '', s)   # strip remaining commands
        return s.strip()

    lines = []
    for part in raw:
        # Split on & (column separator) and clean each column
        cols = [clean(c) for c in part.split('&')]
        lines.append('  '.join(c for c in cols if c))

    return [l for l in lines if l]


# ── XML helpers ────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_colour: str):
    """Set cell background shading (fill) via direct XML."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_colour.lstrip('#'))
    tcPr.append(shd)


def _set_cell_border(cell, top=None, bottom=None, left=None, right=None):
    """Set individual cell borders."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')

    for side, val in (('top', top), ('bottom', bottom),
                      ('left', left), ('right', right)):
        if val is not None:
            el = OxmlElement(f'w:{side}')
            el.set(qn('w:val'),   val.get('val', 'single'))
            el.set(qn('w:sz'),    str(val.get('sz', 4)))
            el.set(qn('w:space'), '0')
            el.set(qn('w:color'), val.get('color', '000000'))
            tcBorders.append(el)

    tcPr.append(tcBorders)


def _remove_cell_border(cell):
    """Remove all borders from a cell (set to 'none')."""
    none_val = {'val': 'none', 'sz': 0, 'color': 'auto'}
    _set_cell_border(cell, none_val, none_val, none_val, none_val)


def _add_paragraph_left_bar(para, hex_colour: str, width_pt: int = 3):
    """Add a coloured left border to a paragraph via pPr/pBdr."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    left.set(qn('w:val'),   'single')
    left.set(qn('w:sz'),    str(width_pt * 8))
    left.set(qn('w:space'), '8')
    left.set(qn('w:color'), hex_colour.lstrip('#'))
    pBdr.append(left)
    pPr.append(pBdr)


# ── Table helpers ──────────────────────────────────────────────────────────────

def _add_calc_row_table(doc: Document, rows: list):
    """
    Add a clean 3-column calc table.
    rows: list of dicts with keys: name, formula, result, label
    """
    if not rows:
        return

    tbl = doc.add_table(rows=len(rows), cols=3)
    tbl.style = 'Table Grid'

    # Column widths: name 3.5 cm | formula 9 cm | result 4 cm
    for i, width in enumerate([Cm(3.5), Cm(9.0), Cm(4.0)]):
        for row in tbl.rows:
            row.cells[i].width = width

    for r_idx, row_data in enumerate(rows):
        cells = tbl.rows[r_idx].cells
        is_alt = r_idx % 2 == 1

        # ── col 0: variable name (bold monospace-ish) ──
        p0 = cells[0].paragraphs[0]
        run0 = p0.add_run(row_data.get('name', ''))
        run0.bold = True
        run0.font.size = Pt(9)
        run0.font.name = 'Consolas'

        # ── col 1: formula ──
        p1 = cells[1].paragraphs[0]
        run1 = p1.add_run(row_data.get('formula', ''))
        run1.font.size = Pt(9)
        run1.font.name = 'Consolas'

        # ── col 2: result (right-aligned, bold) ──
        p2 = cells[2].paragraphs[0]
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run2 = p2.add_run(row_data.get('result', ''))
        run2.bold = True
        run2.font.size = Pt(9)
        run2.font.name = 'Consolas'

        # ── label (annotation) as a second paragraph in col 1 ──
        label = row_data.get('label', '')
        if label:
            p_lbl = cells[1].add_paragraph()
            run_lbl = p_lbl.add_run(f'    {label}')
            run_lbl.italic = True
            run_lbl.font.size = Pt(8)
            run_lbl.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # Subtle alternating background
        bg = 'F7F9FC' if is_alt else 'FFFFFF'
        for cell in cells:
            _set_cell_bg(cell, bg)
            # Remove all heavy borders; keep only a very thin bottom rule
            none_border  = {'val': 'none',   'sz': 0,  'color': 'auto'}
            light_border = {'val': 'single', 'sz': 2,  'color': 'EBEBEB'}
            _set_cell_border(cell,
                top=none_border, bottom=light_border,
                left=none_border, right=none_border)

    # Spacer paragraph after table
    doc.add_paragraph()


# ── Calc-core block renderer ───────────────────────────────────────────────────

def _render_calc_blocks(doc: Document, blocks: list, tmp_files: list):
    """
    Render a list of calc_core blocks (the _result from any calc module)
    into the document.  Consecutive calc_row/handcalc blocks are batched
    into a single table for clean layout.
    """
    # Buffer for consecutive calc_row blocks before flushing to a table
    row_buffer = []

    def flush_rows():
        if row_buffer:
            _add_calc_row_table(doc, list(row_buffer))
            row_buffer.clear()

    for block in blocks:
        btype = block.get('type', '')

        # Silently skip internal sentinel blocks
        if btype.startswith('_'):
            continue

        # ── calc_row → accumulate ──────────────────────────────────────────────
        if btype == 'calc_row':
            row_buffer.append({
                'name':    block.get('name',    ''),
                'formula': block.get('formula', ''),
                'result':  block.get('result',  ''),
                'label':   block.get('label',   ''),
            })
            continue

        # ── handcalc (LaTeX) → parse and accumulate ───────────────────────────
        if btype == 'handcalc':
            flush_rows()  # flush preceding calc_rows first
            lines = _latex_to_lines(block.get('latex', ''))
            for line in lines:
                # Try to split "name = formula = result" pattern
                parts = [p.strip() for p in line.split('=')]
                if len(parts) >= 3:
                    row_buffer.append({
                        'name':    parts[0],
                        'formula': ' = '.join(parts[1:-1]),
                        'result':  parts[-1],
                        'label':   block.get('label', ''),
                    })
                elif len(parts) == 2:
                    row_buffer.append({
                        'name':    parts[0],
                        'formula': '',
                        'result':  parts[1],
                        'label':   block.get('label', ''),
                    })
                else:
                    row_buffer.append({
                        'name': line, 'formula': '', 'result': '', 'label': ''
                    })
            continue

        # Everything else: flush buffered rows first
        flush_rows()

        # ── module_header ──────────────────────────────────────────────────────
        if btype == 'module_header':
            title    = block.get('title',    '')
            subtitle = block.get('subtitle', '')
            material = block.get('material', 'general')
            accent   = _MATERIAL_COLOURS.get(material, _MATERIAL_COLOURS['general'])

            para = doc.add_heading(title, level=1)
            for run in para.runs:
                run.font.color.rgb = _BRAND
            # Coloured left bar
            hex_c = '{:02X}{:02X}{:02X}'.format(*accent)
            _add_paragraph_left_bar(para, hex_c, width_pt=4)

            if subtitle:
                p_sub = doc.add_paragraph()
                run_s = p_sub.add_run(subtitle)
                run_s.italic = True
                run_s.font.size = Pt(9)
                run_s.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        # ── section heading ────────────────────────────────────────────────────
        elif btype == 'section':
            para = doc.add_heading(block.get('content', ''), level=2)
            for run in para.runs:
                run.font.color.rgb = _BRAND

        # ── h1 ────────────────────────────────────────────────────────────────
        elif btype == 'h1':
            doc.add_heading(block.get('content', ''), level=1)

        # ── text ──────────────────────────────────────────────────────────────
        elif btype == 'text':
            content = block.get('content', '').strip()
            if content:
                doc.add_paragraph(content)

        # ── note ──────────────────────────────────────────────────────────────
        elif btype == 'note':
            content = block.get('content', '').strip()
            if content:
                para = doc.add_paragraph()
                run  = para.add_run(f'ℹ  {content}')
                run.italic = True
                run.font.color.rgb = _AMBER
                run.font.size = Pt(9)

        # ── check ─────────────────────────────────────────────────────────────
        elif btype == 'check':
            label  = block.get('label',  '')
            passes = block.get('passes', True)
            value  = block.get('value',  '')
            icon   = '✓' if passes else '✗'
            colour = _GREEN if passes else _RED

            para = doc.add_paragraph()
            run  = para.add_run(f'{icon}  {label}')
            run.bold = True
            run.font.color.rgb = colour
            run.font.size = Pt(9)

            if value:
                run2 = para.add_run(f'   {value}')
                run2.font.size = Pt(9)
                run2.font.color.rgb = colour

        # ── table ─────────────────────────────────────────────────────────────
        elif btype == 'table':
            headers = block.get('headers', [])
            rows    = block.get('rows', [])
            if headers or rows:
                n_cols = max(len(headers), max((len(r) for r in rows), default=0))
                if n_cols == 0:
                    continue
                tbl = doc.add_table(rows=1 + len(rows), cols=n_cols)
                tbl.style = 'Table Grid'

                # Header row
                hdr_cells = tbl.rows[0].cells
                for j, h in enumerate(headers):
                    if j < n_cols:
                        hdr_cells[j].text = str(h)
                        _set_cell_bg(hdr_cells[j], '1E3A5F')
                        for run in hdr_cells[j].paragraphs[0].runs:
                            run.bold = True
                            run.font.color.rgb = RGBColor(0xff, 0xff, 0xff)
                            run.font.size = Pt(9)

                # Data rows
                for r_idx, row_data in enumerate(rows):
                    cells = tbl.rows[r_idx + 1].cells
                    bg = 'F7F9FC' if r_idx % 2 == 1 else 'FFFFFF'
                    for j, cell_val in enumerate(row_data):
                        if j < n_cols:
                            cells[j].text = str(cell_val)
                            _set_cell_bg(cells[j], bg)
                            for run in cells[j].paragraphs[0].runs:
                                run.font.size = Pt(9)

                doc.add_paragraph()

        # ── figure ────────────────────────────────────────────────────────────
        elif btype == 'figure':
            path    = block.get('path', '')
            caption = block.get('caption', '')
            w_mm    = block.get('width_mm', 160)
            if path and os.path.isfile(path):
                try:
                    doc.add_picture(path, width=Cm(w_mm / 10.0))
                    if caption:
                        cap_p = doc.add_paragraph(caption)
                        cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        for run in cap_p.runs:
                            run.italic = True
                            run.font.size = Pt(8)
                except Exception as exc:
                    doc.add_paragraph(f'[Figure could not be embedded: {exc}]')

    # Flush any remaining calc_row buffer
    flush_rows()


# ── Document-level block converters ───────────────────────────────────────────

_CALC_TYPES = {
    "steel_beam", "steel_column", "beam_column",
    "rc_beam", "rc_column", "rc_slab",
    "timber_beam", "timber_column",
    "masonry_wall",
    "custom_calc",
    "wind_load", "snow_load",
    "foundation",
    "load_combo",
    "bolt_group", "fillet_weld",
    "plate_girder",
    "saved_calc",
}


def _convert_doc_block(doc: Document, block: dict, tmp_files: list):
    btype = block.get('type', '')

    # ── heading ───────────────────────────────────────────────────────────────
    if btype == 'heading':
        level = block["data"].get("level", 1)
        text  = block["data"].get("text", "")
        level = max(1, min(3, level))
        doc.add_heading(text, level=level)
        return

    # ── text ──────────────────────────────────────────────────────────────────
    if btype == 'text':
        text = block["data"].get("text", "").strip()
        if text:
            doc.add_paragraph(text)
        return

    # ── image ─────────────────────────────────────────────────────────────────
    if btype == 'image':
        b64     = block["data"].get("image_b64", "")
        caption = block["data"].get("caption", "")
        w_pct   = block["data"].get("width_pct", 100)
        w_cm    = max(4.0, min(17.0, 17.0 * w_pct / 100.0))
        if b64:
            try:
                header, data = b64.split(",", 1)
                ext = "jpg" if ("jpeg" in header or "jpg" in header) else "png"
                tmp = tempfile.NamedTemporaryFile(suffix=f".{ext}", delete=False)
                tmp.write(base64.b64decode(data))
                tmp.close()
                tmp_files.append(tmp.name)
                doc.add_picture(tmp.name, width=Cm(w_cm))
                if caption:
                    p = doc.add_paragraph(caption)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.italic = True
                        run.font.size = Pt(8)
            except Exception as exc:
                doc.add_paragraph(f'[Image could not be embedded: {exc}]')
        return

    # ── python_calc ────────────────────────────────────────────────────────────
    if btype == 'python_calc':
        d = block["data"]
        doc.add_heading(d.get("title", "Python Script"), level=2)
        output = d.get("_output_text", "").strip()
        if output:
            para = doc.add_paragraph()
            run  = para.add_run(output)
            run.font.name = 'Consolas'
            run.font.size = Pt(8)
        if d.get("_error", "").strip():
            doc.add_paragraph(f'Script error: {d["_error"]}')
        # Figures from python blocks
        for b64 in d.get("_figs_b64", []):
            try:
                tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                tmp.write(base64.b64decode(b64))
                tmp.close()
                tmp_files.append(tmp.name)
                doc.add_picture(tmp.name, width=Cm(14))
            except Exception:
                pass
        return

    # ── beam_fem ───────────────────────────────────────────────────────────────
    if btype == 'beam_fem':
        d      = block["data"]
        result = d.get("_result")
        if result is None:
            doc.add_heading(d.get("title", "Beam FEM Analysis"), level=2)
            doc.add_paragraph("This block has not been run yet.")
            return
        _render_calc_blocks(doc, result, tmp_files)
        # Embed FEM figure
        fig_b64 = d.get("_fig_b64")
        if fig_b64:
            try:
                tmp = tempfile.NamedTemporaryFile(suffix=".png", delete=False)
                tmp.write(base64.b64decode(fig_b64))
                tmp.close()
                tmp_files.append(tmp.name)
                doc.add_picture(tmp.name, width=Cm(16))
            except Exception:
                pass
        return

    # ── all calc types ─────────────────────────────────────────────────────────
    if btype in _CALC_TYPES:
        d      = block["data"]
        result = d.get("_result")
        if result is None:
            title = d.get("title", btype.replace("_", " ").title())
            doc.add_heading(title, level=2)
            doc.add_paragraph("This block has not been run yet — open the editor and click 'Run check'.")
            return
        _render_calc_blocks(doc, result, tmp_files)
        return

    # Unknown block types — skip silently


# ── Heading numbering (mirror of pdf_builder) ─────────────────────────────────

def _number_headings(blocks: list) -> list:
    result   = []
    counters = [0, 0, 0]
    for block in blocks:
        if block.get("type") == "heading":
            level = block["data"].get("level", 1)
            text  = block["data"].get("text", "")
            idx   = level - 1
            if idx < 3:
                counters[idx] += 1
                for j in range(idx + 1, 3):
                    counters[j] = 0
                num      = ".".join(str(counters[k]) for k in range(idx + 1))
                new_text = f"{num}  {text}" if text else num
                block = {**block, "data": {**block["data"], "text": new_text}}
        result.append(block)
    return result


# ── Cover page ────────────────────────────────────────────────────────────────

def _add_cover(doc: Document, project: dict, doc_id: str):
    """Insert a minimal title page at the beginning of the document."""
    meta = project.get("metadata", {})

    # Large title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p_title.add_run(meta.get("project_name", "Structural Calculations"))
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = _BRAND

    # Subtitle line
    sub = meta.get("project_ref", "") or doc_id
    if sub:
        p_sub = doc.add_paragraph()
        p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p_sub.add_run(sub)
        run2.font.size = Pt(13)
        run2.font.color.rgb = RGBColor(0x44, 0x55, 0x66)

    doc.add_paragraph()   # blank line

    # Meta table (2 columns)
    rows = []
    for key, label in [
        ("engineer",  "Engineer"),
        ("checker",   "Checker"),
        ("client",    "Client"),
        ("date",      "Date"),
        ("revision",  "Revision"),
        ("standard",  "Standard"),
    ]:
        val = meta.get(key, "")
        if val:
            rows.append((label, val))

    if rows:
        tbl = doc.add_table(rows=len(rows), cols=2)
        for i, (label, value) in enumerate(rows):
            cells = tbl.rows[i].cells
            cells[0].text = label
            cells[1].text = value
            for run in cells[0].paragraphs[0].runs:
                run.bold = True
                run.font.size = Pt(10)
            for run in cells[1].paragraphs[0].runs:
                run.font.size = Pt(10)
            # Remove cell borders
            for cell in cells:
                _remove_cell_border(cell)

    # Page break after cover
    doc.add_page_break()


# ── Document styles ────────────────────────────────────────────────────────────

def _apply_styles(doc: Document):
    """
    Tune the built-in Normal / Heading styles for a clean engineering look.
    Operates on whatever styles are available in the default docx template.
    """
    from docx.shared import Pt

    # Normal text
    normal = doc.styles['Normal']
    normal.font.name = 'Calibri'
    normal.font.size = Pt(10)

    # Heading 1
    try:
        h1 = doc.styles['Heading 1']
        h1.font.name = 'Calibri'
        h1.font.size = Pt(14)
        h1.font.bold = True
        h1.font.color.rgb = _BRAND
    except Exception:
        pass

    # Heading 2
    try:
        h2 = doc.styles['Heading 2']
        h2.font.name = 'Calibri'
        h2.font.size = Pt(12)
        h2.font.bold = True
        h2.font.color.rgb = _BRAND
    except Exception:
        pass

    # Heading 3
    try:
        h3 = doc.styles['Heading 3']
        h3.font.name = 'Calibri'
        h3.font.size = Pt(11)
        h3.font.bold = True
        h3.font.color.rgb = RGBColor(0x33, 0x55, 0x77)
    except Exception:
        pass


# ── Main entry point ───────────────────────────────────────────────────────────

def build_word(project: dict, blocks: list, doc_id: str = "") -> bytes:
    """
    Convert a list of v2 document blocks to a .docx and return the bytes.

    Parameters
    ----------
    project : dict   — Full project object (metadata used for cover page).
    blocks  : list   — Blocks from project["documents"][doc_id]["blocks"].
    doc_id  : str    — Document identifier (shown on cover page subtitle).
    """
    tmp_files: list = []

    doc = Document()
    _apply_styles(doc)

    # Cover page
    _add_cover(doc, project, doc_id)

    # Number headings
    numbered_blocks = _number_headings(blocks)

    for block in numbered_blocks:
        try:
            _convert_doc_block(doc, block, tmp_files)
        except Exception as exc:
            doc.add_paragraph(f"[Error rendering block '{block.get('type')}': {exc}]")

    # Return as bytes
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    docx_bytes = buf.read()

    # Clean up temp image files
    for path in tmp_files:
        try:
            os.unlink(path)
        except Exception:
            pass

    return docx_bytes
